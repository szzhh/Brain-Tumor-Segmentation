import numpy as np
import nibabel as nib
import cv2
import os
import requests
import json
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
from PySide2.QtGui import QImage
import matplotlib.pyplot as plt


def gz2np(path_mod,path_seg):

    lst1=[]
    lst2=[]
    seg = nib.load(path_seg)
    mod_t1 = nib.load(path_mod)

    seg_data = np.array(seg.dataobj)
    mod_t1_data = np.array(mod_t1.dataobj)
    r = np.array([255, 0, 0]).squeeze() 
    g = np.array([0, 255, 0]).squeeze()
    b = np.array([0, 0, 255]).squeeze()

    for n in range(155):
        temp_img = mod_t1_data[:, :, n]
        img = temp_img/3
        img = np.expand_dims(img, axis=2)
        img = np.concatenate((img, img, img), axis=-1)
        lst1.append(img)
        
        rgb_img = img.copy()
        pos = np.where(seg_data[:, :, n] == 1)
        rgb_img[pos] = r
        pos = np.where(seg_data[:, :, n] == 2)
        rgb_img[pos] = g
        pos = np.where(seg_data[:, :, n] == 4)
        rgb_img[pos] = b
        lst2.append(rgb_img)
        
    lst1=np.array(lst1)
    lst2=np.array(lst2)
    lst1=lst1.astype(np.uint8)
    lst2=lst2.astype(np.uint8)
    return lst1,lst2

def cov(imgg):
    imgg = QImage(imgg[:], imgg.shape[1], imgg.shape[0], imgg.shape[1] * 3, QImage.Format_RGB888)
    return imgg

def req_get(puid,modtype):
    if not os.path.exists(os.path.split(os.path.realpath(__file__))[0]+r'/temp/'): 
        os.mkdir(os.path.split(os.path.realpath(__file__))[0]+r'/temp/')
    r = requests.get('http://172.24.21.239:5000/download/'+puid+'/'+modtype)
    with open(os.path.split(os.path.realpath(__file__))[0]+r'/temp/'+modtype+'.nii.gz', 'wb') as f:
        f.write(r.content)
    return os.path.split(os.path.realpath(__file__))[0]+r'/temp/'+modtype+'.nii.gz'

def generatePDF(modulePath, savePath, imagesPath, advicePath, V1,V2,V3, kwargs):
    """
    describe: generate PDF from given data and save to specific path.
    params: 
        modulePath: String. Module path.
        savePath: String. Save path.
        imagesPath: Tuple [img1, img2]. Images' path.
        advicePath: String. Prior knowledge in treating tumor. A json formated file.
        TumorR: R of tumor. 
        TumorArea: Float. The area of tumor.
        **kwargs: Dict. Information of the patient. Keys: ['姓名', '性别', '年龄', '编号', '住院号', '床号']. Value: String.
    """
    with open(advicePath, 'r',encoding='utf-8') as f:
        infoDict = json.load(f)

    doc = Document(modulePath)

    for i, para in enumerate(doc.paragraphs):
        if i == 1:
            # para.text = '姓名：{:^10} 性别：{:^10} 年龄：{:^10} 申请科室：{:^10}'.format('张三', '男', '20', '肾脏科室')
            para.style = doc.styles['Normal']
            para.style.font.size = 143350
            # 增添姓名
            para.add_run('姓名：').bold = True
            para.add_run('{0:{1}^8}'.format(kwargs['姓名'], chr(12288)))
            # 增添性别
            para.add_run('性别：').bold = True
            para.add_run('{0:{1}^8}'.format(kwargs['性别'], chr(12288)))
            # 增添年龄
            para.add_run('年龄：').bold = True
            para.add_run('{:^15}'.format(kwargs['年龄']))
            # 增添科室
            para.add_run('申请科室：').bold = True
            para.add_run('{0:{1}^8}'.format('肿瘤内科', chr(12288)))
        
        elif i == 2:
            para.style = doc.styles['Normal']
            para.style.font.size = 143350
            # 增添姓名
            para.add_run('编号：').bold = True
            para.add_run('{:^16}'.format(kwargs['编号']))
            # 增添性别
            para.add_run('住院号：').bold = True
            para.add_run('{:^14}'.format(kwargs['住院号']))
            # 增添年龄
            para.add_run('床号：').bold = True
            para.add_run('{:^15}'.format(kwargs['床号']))
            # 增添科室
            para.add_run('检查部位：').bold = True
            para.add_run('{:^15}'.format('脑部'))

        elif i == 4:
            para.text = ' ' * 90
            para.runs[-1].add_picture(imagesPath[0], width=Inches(3.25))
            para.add_run('{:^2}'.format(' '))
            para.runs[-1].add_picture(imagesPath[1], width=Inches(3.23))
        
        elif i == 8:
            descriptition = '发现脑肿瘤，肿瘤体积为{}mm^3，其核心肿瘤体积为{}mm^3，其增强肿瘤体积为{}mm^3。'.format(V1,V2,V3)
            stageMap = {
                '1': 'I期脑细胞癌',
                '2': 'II期脑细胞癌',
                '3': 'III期脑细胞癌',
                '4': 'IV期脑细胞癌'
            }
            # if TumorR <= 7:
            #     stage = '1'
            # elif TumorR > 7:
            stage = '2'
            descriptition += '经估计可能为{}。'.format(stageMap[stage])
            para.style = doc.styles['Normal']
            para.style.font.size = 143350
            para.add_run('\t'+descriptition)  
            
        
        elif i == 10:
            para.style = doc.styles['Normal']
            para.style.font.size = 143350
            advice = infoDict[str(stage)]
            para.add_run('\t'+advice[0]).bold = True
            para.add_run('\n')
            for ad in advice[1:]:
                para.add_run('\t\t\t'+ad)
                para.add_run('\n')

    docp = savePath+'/{}-{}.docx'.format(kwargs['编号'], kwargs['姓名'])
    doc.save(docp)
    pdfpath=docp[:-4]+'pdf'
    convert(docp, pdfpath)
    os.remove(docp)
    return pdfpath

if __name__ == '__main__':
    lst1,lst2=gz2np(os.path.split(os.path.realpath(__file__))[0]+'/temp/t1.nii.gz',os.path.split(os.path.realpath(__file__))[0]+'/temp/seg.nii.gz')
    # np.save(os.path.split(os.path.realpath(__file__))[0]+'/output/lst1.npy',lst1)
    # np.save(os.path.split(os.path.realpath(__file__))[0]+'/output/lst2.npy',lst2)
    plt.imshow(lst1[101])
    plt.show()
    plt.imshow(lst2[101])
    plt.show()